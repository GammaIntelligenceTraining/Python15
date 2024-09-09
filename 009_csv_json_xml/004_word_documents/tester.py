from docx import Document
import docx

document = Document('tester.docx')
print(len(document.paragraphs))
print(document.paragraphs[0].text)
print(document.paragraphs[0].runs)


def get_text(filename):
    document = docx.Document(filename)
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)
    return full_text

print(get_text('tester.docx'))